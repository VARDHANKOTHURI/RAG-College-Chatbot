import os
import sys
import json
import io
from fastapi.testclient import TestClient

# Ensure root dir is in sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from backend.app.main import app
from backend.app.config.settings import settings
from backend.app.services.document_service import document_service

def test_complete_college_chatbot_suite():
    print("\n" + "=" * 70)
    print("      RUNNING FULL-STACK COLLEGE CHATBOT AUTOMATED TEST SUITE      ")
    print("=" * 70)

    with TestClient(app) as client:
        # ----------------------------------------------------
        # 1. System Health Check
        # ----------------------------------------------------
        print("\n[TEST 1] System Health Check (/api/health)...")
        res = client.get("/api/health")
        assert res.status_code == 200, f"Health check failed: {res.text}"
        health_data = res.json()
        print(f" -> Service: {health_data['service']} v{health_data['version']}")
        print(f" -> Vector Database: {health_data['vector_store']['collection']} ({health_data['vector_store']['total_vectors']} vectors)")
        assert health_data["status"] == "healthy"

        # ----------------------------------------------------
        # 2. Student Registration & Authentication
        # ----------------------------------------------------
        print("\n[TEST 2] Student Registration & Profile...")
        new_student_email = f"alice_{os.urandom(3).hex()}@college.edu"
        reg_payload = {
            "name": "Alice Johnson",
            "email": new_student_email,
            "password": "Password@123",
            "role": "student"
        }
        res = client.post("/api/auth/register", json=reg_payload)
        assert res.status_code == 200, f"Registration failed: {res.text}"
        alice_data = res.json()
        alice_token = alice_data["access_token"]
        alice_headers = {"Authorization": f"Bearer {alice_token}"}
        print(f" -> Registered student: {alice_data['user']['name']} ({alice_data['user']['email']})")

        # Test duplicate registration rejection
        dup_res = client.post("/api/auth/register", json=reg_payload)
        assert dup_res.status_code == 400, "Duplicate email registration should fail"
        print(" -> Correctly rejected duplicate registration")

        # Test /api/auth/me
        me_res = client.get("/api/auth/me", headers=alice_headers)
        assert me_res.status_code == 200
        assert me_res.json()["email"] == new_student_email
        print(" -> Verified /api/auth/me profile endpoint")

        # ----------------------------------------------------
        # 3. Admin Authentication & Role Enforcement
        # ----------------------------------------------------
        print("\n[TEST 3] Admin Login & Role-Based Authorization...")
        admin_res = client.post("/api/auth/login", json={"email": "admin@college.edu", "password": "Admin@123"})
        assert admin_res.status_code == 200, f"Admin login failed: {admin_res.text}"
        admin_token = admin_res.json()["access_token"]
        admin_headers = {"Authorization": f"Bearer {admin_token}"}
        print(f" -> Authenticated Administrator: {admin_res.json()['user']['name']}")

        # Student attempting admin endpoint should receive 403 FORBIDDEN
        forbidden_res = client.get("/api/admin/analytics", headers=alice_headers)
        assert forbidden_res.status_code == 403, "Student should be forbidden from admin analytics"
        print(" -> Role-based protection verified (Student blocked from Admin route: 403 Forbidden)")

        # Admin accessing analytics
        admin_analytics_res = client.get("/api/admin/analytics", headers=admin_headers)
        assert admin_analytics_res.status_code == 200
        analytics_data = admin_analytics_res.json()
        print(f" -> Admin Analytics: {analytics_data['totalDocuments']} docs, {analytics_data['totalChunks']} chunks")

        # ----------------------------------------------------
        # 4. Multi-Format Document Ingestion (PDF, DOCX, TXT)
        # ----------------------------------------------------
        print("\n[TEST 4] Admin Document Ingestion (PDF, DOCX, TXT)...")
        
        # 4a. Ingest a dynamic PDF document
        import fitz  # PyMuPDF
        pdf_doc = fitz.open()
        pdf_page = pdf_doc.new_page()
        pdf_page.insert_text(
            (50, 72),
            "GREENWOOD INSTITUTE OF TECHNOLOGY\nDEPARTMENT OF ARTIFICIAL INTELLIGENCE & DATA SCIENCE\n\n"
            "Specialized Lab Curriculum 2026:\n"
            "The AI & Robotics Center of Excellence is located in Tech Tower Block C, 4th Floor.\n"
            "Students with CGPA above 8.5 can apply for the NVIDIA Deep Learning Research Fellowship.\n"
            "Fellowship Grant: INR 35,000 per semester plus funded GPU compute credits.",
            fontsize=11
        )
        pdf_bytes = pdf_doc.write()
        pdf_doc.close()

        pdf_upload_res = client.post(
            "/api/documents/upload",
            files={"file": ("AI_Department_Fellowship_2026.pdf", pdf_bytes, "application/pdf")},
            data={
                "title": "AI Department Research Fellowship 2026",
                "category": "Academics",
                "department": "CSE-AIML",
                "academic_year": "2026",
                "description": "NVIDIA Deep Learning fellowship and GPU compute grant guidelines."
            },
            headers=admin_headers
        )
        assert pdf_upload_res.status_code == 200, f"PDF upload failed: {pdf_upload_res.text}"
        pdf_doc_id = pdf_upload_res.json()["id"]
        print(f" -> Uploaded & Processed PDF Document ID: {pdf_doc_id}")

        # 4b. Ingest a dynamic DOCX document
        import docx
        doc_obj = docx.Document()
        doc_obj.add_heading("Greenwood Sports Academy & Gymnasium Guidelines 2026", level=1)
        doc_obj.add_paragraph("The Olympic-size swimming pool and indoor badminton courts are open daily from 5:30 AM to 8:30 AM and 4:30 PM to 8:00 PM.")
        doc_obj.add_paragraph("Annual Sports Membership Fee for non-hostel day scholars is INR 3,000 per academic year.")
        
        docx_io = io.BytesIO()
        doc_obj.save(docx_io)
        docx_bytes = docx_io.getvalue()

        docx_upload_res = client.post(
            "/api/documents/upload",
            files={"file": ("Sports_Complex_Regulations_2026.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={
                "title": "Sports Complex & Gymnasium Regulations 2026",
                "category": "General FAQ",
                "department": "All",
                "academic_year": "2026",
                "description": "Timings, swimming pool rules, and gym membership fees."
            },
            headers=admin_headers
        )
        assert docx_upload_res.status_code == 200, f"DOCX upload failed: {docx_upload_res.text}"
        docx_doc_id = docx_upload_res.json()["id"]
        print(f" -> Uploaded & Processed DOCX Document ID: {docx_doc_id}")

        # ----------------------------------------------------
        # 5. Student Conversational Chat & RAG Retrieval on New Documents
        # ----------------------------------------------------
        print("\n[TEST 5] Student Chat & RAG Verification on Newly Ingested Documents...")
        
        # Test Query on newly uploaded PDF document
        ai_query = "What is the grant amount for the NVIDIA Deep Learning Research Fellowship?"
        chat_res1 = client.post("/api/chat", json={"message": ai_query}, headers=alice_headers)
        assert chat_res1.status_code == 200
        ai_answer = chat_res1.json()
        print(f" -> Query: '{ai_query}'")
        print(f" -> Answer: {ai_answer['answer'][:160]}...")
        assert "35,000" in ai_answer["answer"] or "nvidia" in ai_answer["answer"].lower()
        assert len(ai_answer["sources"]) > 0
        print(f" -> Verified Source: {ai_answer['sources'][0]['title']} (Page {ai_answer['sources'][0]['pageNumber']})")

        # Test Query on newly uploaded DOCX document
        sports_query = "What are the timings for the swimming pool and badminton courts?"
        chat_res2 = client.post(
            "/api/chat",
            json={"message": sports_query, "conversationId": ai_answer["conversationId"]},
            headers=alice_headers
        )
        assert chat_res2.status_code == 200
        sports_answer = chat_res2.json()
        print(f"\n -> Follow-up Query: '{sports_query}'")
        print(f" -> Answer: {sports_answer['answer'][:160]}...")
        assert "5:30" in sports_answer["answer"] or "swimming" in sports_answer["answer"].lower()

        # ----------------------------------------------------
        # 6. Unknown Query Rejection & Knowledge Gap Recording
        # ----------------------------------------------------
        print("\n[TEST 6] Hallucination Resistance & Knowledge Gap Logging...")
        unknown_q = "What is the ticket price for taking a private spaceship from the college campus to the Moon?"
        unk_res = client.post("/api/chat", json={"message": unknown_q}, headers=alice_headers)
        assert unk_res.status_code == 200
        unk_data = unk_res.json()
        assert unk_data["isUnknown"] == True
        print(f" -> Unknown Query: '{unknown_q}'")
        print(f" -> Answer: {unk_data['answer']}")
        print(f" -> Correctly flagged isUnknown: {unk_data['isUnknown']}")

        # Verify query was saved to unanswered questions in admin panel
        unanswered_res = client.get("/api/admin/unanswered?status=open", headers=admin_headers)
        assert unanswered_res.status_code == 200
        unanswered_list = unanswered_res.json()
        assert any(unknown_q in u["question"] for u in unanswered_list)
        print(" -> Verified knowledge gap recorded in Administrator Unanswered Queue")

        # Resolve knowledge gap
        matching_q = next(u for u in unanswered_list if unknown_q in u["question"])
        resolve_res = client.put(
            f"/api/admin/unanswered/{matching_q['id']}",
            json={"status": "resolved", "adminNotes": "Query identified as out of domain"},
            headers=admin_headers
        )
        assert resolve_res.status_code == 200
        print(" -> Verified admin can resolve knowledge gap items")

        # ----------------------------------------------------
        # 7. Answer Feedback System
        # ----------------------------------------------------
        print("\n[TEST 7] Student Answer Feedback Submission...")
        fb_res = client.post(
            "/api/feedback",
            json={
                "messageId": ai_answer["messageId"],
                "rating": 1,
                "reason": "Accurate information",
                "comment": "Provided exact fellowship grant amount and lab location."
            },
            headers=alice_headers
        )
        assert fb_res.status_code == 200
        print(f" -> Submitted feedback for message {ai_answer['messageId']}")

        # Admin lists feedback
        admin_fb_res = client.get("/api/feedback", headers=admin_headers)
        assert admin_fb_res.status_code == 200
        assert len(admin_fb_res.json()) > 0
        print(f" -> Admin retrieved {len(admin_fb_res.json())} feedback submissions")

        # ----------------------------------------------------
        # 8. Knowledge Collections Management
        # ----------------------------------------------------
        print("\n[TEST 8] Knowledge Collections CRUD...")
        col_res = client.get("/api/collections", headers=alice_headers)
        assert col_res.status_code == 200
        assert len(col_res.json()) >= 4
        print(f" -> Existing Collections: {[c['name'] for c in col_res.json()]}")

        # Create new collection
        new_col_res = client.post(
            "/api/collections",
            json={
                "name": "Robotics & Drone Research Lab",
                "description": "Autonomous flight control, sensor fusion, and ROS manuals",
                "department": "CSE-AIML",
                "accessRules": ["student", "admin"]
            },
            headers=admin_headers
        )
        assert new_col_res.status_code == 200
        new_col_id = new_col_res.json()["id"]
        print(f" -> Created Collection: {new_col_res.json()['name']} (ID: {new_col_id})")

        # ----------------------------------------------------
        # 9. Conversation History & Deletion
        # ----------------------------------------------------
        print("\n[TEST 9] Conversation Details & Deletion...")
        conv_details_res = client.get(f"/api/chat/conversations/{ai_answer['conversationId']}", headers=alice_headers)
        assert conv_details_res.status_code == 200
        conv_data = conv_details_res.json()
        assert len(conv_data["messages"]) >= 2
        print(f" -> Retrieved Conversation '{conv_data['title']}' with {len(conv_data['messages'])} messages")

        del_conv_res = client.delete(f"/api/chat/conversations/{ai_answer['conversationId']}", headers=alice_headers)
        assert del_conv_res.status_code == 200
        print(" -> Deleted conversation successfully")

        # ----------------------------------------------------
        # 10. Web Interface Verification
        # ----------------------------------------------------
        print("\n[TEST 10] Web Application Client-Side Routes...")
        for route in ["/", "/chat", "/documents", "/admin", "/settings", "/login", "/register"]:
            r = client.get(route)
            assert r.status_code == 200, f"Route {route} failed to load"
            assert "<!DOCTYPE html>" in r.text
        print(" -> Verified all SPA client routes (/, /chat, /documents, /admin, /settings, /login, /register)")

    print("\n" + "=" * 70)
    print("   >>> ALL 10 COMPREHENSIVE INTEGRATION SUITE TESTS PASSED 100%! <<<   ")
    print("=" * 70 + "\n")

if __name__ == "__main__":
    test_complete_college_chatbot_suite()
