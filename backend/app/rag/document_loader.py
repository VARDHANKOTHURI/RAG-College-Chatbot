import os
from typing import List, Dict, Any
from backend.app.utils.logger import logger

class LoadedPage:
    def __init__(self, page_number: int, text: str, metadata: Dict[str, Any] = None):
        self.page_number = page_number
        self.text = text
        self.metadata = metadata or {}

class DocumentLoader:
    @staticmethod
    def load_pdf(file_path: str) -> List[LoadedPage]:
        pages = []
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            for page_idx in range(len(doc)):
                page = doc.load_page(page_idx)
                text = page.get_text("text")
                pages.append(LoadedPage(page_number=page_idx + 1, text=text, metadata={"total_pages": len(doc)}))
            doc.close()
        except ImportError:
            logger.warning("PyMuPDF (fitz) not installed. Trying pypdf/fallback.")
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                for idx, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    pages.append(LoadedPage(page_number=idx + 1, text=text, metadata={"total_pages": len(reader.pages)}))
            except Exception as e:
                logger.error(f"Failed to extract PDF with pypdf fallback: {e}")
                raise
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
        return pages

    @staticmethod
    def load_docx(file_path: str) -> List[LoadedPage]:
        pages = []
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        full_text.append(row_text)
            content = "\n".join(full_text)
            pages.append(LoadedPage(page_number=1, text=content, metadata={"total_pages": 1}))
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
        return pages

    @staticmethod
    def load_txt(file_path: str) -> List[LoadedPage]:
        encodings = ["utf-8", "latin-1", "windows-1252"]
        content = ""
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        # Split txt into virtual pages if large or containing page markers
        if "--- Page " in content or "=== Page " in content:
            parts = content.split("--- Page ")
            pages = []
            for idx, part in enumerate(parts):
                if part.strip():
                    pages.append(LoadedPage(page_number=idx + 1, text=part, metadata={"total_pages": len(parts)}))
            return pages
        else:
            return [LoadedPage(page_number=1, text=content, metadata={"total_pages": 1})]

    @classmethod
    def load(cls, file_path: str) -> List[LoadedPage]:
        ext = file_path.split(".")[-1].lower()
        if ext == "pdf":
            return cls.load_pdf(file_path)
        elif ext in ["docx", "doc"]:
            return cls.load_docx(file_path)
        elif ext in ["txt", "md", "csv"]:
            return cls.load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: .{ext}")
